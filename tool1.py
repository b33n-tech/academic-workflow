import streamlit as st
import uuid
import json
from docx import Document
from io import BytesIO
import streamlit.components.v1 as components

st.set_page_config(page_title='Tool 1 - Architecture Builder', layout='wide')

# --- Helpers ---
LEVELS = {
    'Titre 1: I.': 1,
    'Titre 2: I.1.': 2,
    'Titre 3: I.1.a.': 3,
    'Titre 4: I.1.a.1.': 4,
    'Titre 5: I.1.a.1.a': 5,
}
LLM_OPTIONS = [
    'https://claude.ai/',
    'https://chat.deepseek.com/',
    'https://chatgpt.com/',
    'https://www.perplexity.ai/'
]

if 'blocks' not in st.session_state:
    st.session_state.blocks = []
if 'history' not in st.session_state:
    st.session_state.history = []
if 'selected' not in st.session_state:
    st.session_state.selected = None


def push_history():
    st.session_state.history.append(json.dumps(st.session_state.blocks))
    # keep history reasonable size
    if len(st.session_state.history) > 50:
        st.session_state.history.pop(0)


def undo():
    if st.session_state.history:
        last = st.session_state.history.pop()
        st.session_state.blocks = json.loads(last)


def new_block(title='Untitled', level='Titre 2: I.1.'):
    return {
        'id': str(uuid.uuid4()),
        'title': title,
        'level': level,
        'comments': [],
        'pitch': '',
        'summary': ''
    }


def export_json():
    return json.dumps(st.session_state.blocks, ensure_ascii=False, indent=2)


def export_docx():
    doc = Document()
    for b in st.session_state.blocks:
        lvl = LEVELS.get(b['level'], 1)
        # map level to heading style (1-5)
        heading = min(lvl, 4)
        doc.add_heading(b['title'] or 'Untitled', level=heading)
        # include short pitch if exists
        if b.get('pitch'):
            doc.add_paragraph('Pitch: ' + b['pitch'])
        if b.get('summary'):
            doc.add_paragraph(b['summary'])
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- Top controls ---
col1, col2, col3, col4 = st.columns([1,1,6,2])
with col1:
    if st.button('+ Ajouter un bloc'):
        push_history()
        st.session_state.blocks.append(new_block())
with col2:
    if st.button('- Supprimer le bloc sélectionné'):
        if st.session_state.selected:
            push_history()
            st.session_state.blocks = [b for b in st.session_state.blocks if b['id'] != st.session_state.selected]
            st.session_state.selected = None
with col3:
    st.write('')
with col4:
    st.session_state.selected_llm = st.selectbox('LLM', LLM_OPTIONS, index=0)

st.markdown('---')

# Upload archive
uploaded = st.file_uploader('Charger une archive JSON', type=['json'])
if uploaded is not None:
    try:
        data = json.load(uploaded)
        if isinstance(data, list):
            push_history()
            st.session_state.blocks = data
            st.success('Archive chargée ✅')
        else:
            st.error('Le JSON doit contenir une liste de blocs')
    except Exception as e:
        st.error('Erreur parsing JSON: ' + str(e))

# --- Main area: blocks list ---
st.write('### Architecture (glissez via boutons "↑"/"↓")')
for i, b in enumerate(st.session_state.blocks):
    cols = st.columns([1,6,1,1,1])
    with cols[0]:
        if st.button('↑', key=f'up_{b["id"]}') and i > 0:
            push_history()
            st.session_state.blocks[i], st.session_state.blocks[i-1] = st.session_state.blocks[i-1], st.session_state.blocks[i]
    with cols[1]:
        expanded = st.expander(f"{b.get('level')} {b.get('title') or 'Untitled'}", expanded=False)
        with expanded:
            title = st.text_input('Titre', value=b['title'], key=f'title_{b["id"]}')
            level = st.selectbox('Niveau', list(LEVELS.keys()), index=list(LEVELS.keys()).index(b['level']), key=f'level_{b["id"]}')
            st.session_state.blocks[i]['title'] = title
            st.session_state.blocks[i]['level'] = level
            st.session_state.selected = b['id']
            st.write('Commentaires:')
            for ci, c in enumerate(b.get('comments', [])):
                st.write(f"- [{c['ts']}] {c['text']}")
    with cols[2]:
        if st.button('↓', key=f'down_{b["id"]}') and i < len(st.session_state.blocks)-1:
            push_history()
            st.session_state.blocks[i], st.session_state.blocks[i+1] = st.session_state.blocks[i+1], st.session_state.blocks[i]
    with cols[3]:
        if st.button('Suppr', key=f'del_{b["id"]}'):
            push_history()
            st.session_state.blocks.pop(i)
            st.experimental_rerun()
    with cols[4]:
        if st.button('Sélectionner', key=f'sel_{b["id"]}'):
            st.session_state.selected = b['id']

st.markdown('---')

# --- Floating pad (simulated in sidebar) ---
st.sidebar.title('Pad flottant')
if st.sidebar.button('Ajouter un commentaire'):
    if not st.session_state.selected:
        st.sidebar.warning('Sélectionne d\'abord un bloc')
    else:
        text = st.sidebar.text_area('Texte du commentaire', key='new_comment')
        if st.sidebar.button('Valider commentaire'):
            push_history()
            for b in st.session_state.blocks:
                if b['id'] == st.session_state.selected:
                    b.setdefault('comments', []).append({'ts': st.session_state.get('timer', 'now'), 'text': text})
                    st.sidebar.success('Commentaire ajouté')

pitch = st.sidebar.text_input('Ajouter un pitch/sujet rapide', key='pad_pitch')
if st.sidebar.button('Valider pitch'):
    if not st.session_state.selected:
        st.sidebar.warning('Sélectionne d\'abord un bloc')
    else:
        push_history()
        for b in st.session_state.blocks:
            if b['id'] == st.session_state.selected:
                b['pitch'] = pitch
                st.sidebar.success('Pitch ajouté')

summary = st.sidebar.text_area('Ajouter un résumé long', key='pad_summary')
if st.sidebar.button('Valider résumé'):
    if not st.session_state.selected:
        st.sidebar.warning('Sélectionne d\'abord un bloc')
    else:
        push_history()
        for b in st.session_state.blocks:
            if b['id'] == st.session_state.selected:
                b['summary'] = summary
                st.sidebar.success('Résumé ajouté')

if st.sidebar.button('Télécharger archive JSON'):
    st.sidebar.download_button('Télécharger JSON', data=export_json(), file_name='architecture.json')

# --- Bottom buttons ---
c1, c2, c3, c4 = st.columns(4)
with c1:
    st.download_button('Télécharger JSON', data=export_json(), file_name='architecture.json')
with c2:
    docx_bio = None
    if st.button('Générer DOCX'): 
        docx_bio = export_docx()
        st.download_button('Télécharger DOCX', data=docx_bio, file_name='architecture.docx')
with c3:
    if st.button('Undo'):
        undo()
with c4:
    if st.button('Push LLM (copier titres hiérarchisés)'):
        # build hierarchical titles
        lines = []
        counters = {1:0,2:0,3:0,4:0,5:0}
        for b in st.session_state.blocks:
            lvl = LEVELS.get(b['level'],1)
            counters[lvl] += 1
            # reset deeper counters
            for k in range(lvl+1,6):
                counters[k]=0
            # build numbering like I., I.1., I.1.a. (simplified using numbers/letters)
            numbering = ''
            if lvl >= 1:
                numbering += str(counters[1]) + '.'
            if lvl >= 2:
                numbering += str(counters[2]) + '.'
            if lvl >= 3:
                numbering += chr(96 + counters[3]) + '.' if counters[3]>0 else ''
            if lvl >= 4:
                numbering += str(counters[4]) + '.'
            if lvl >= 5:
                numbering += chr(96 + counters[5]) + '.' if counters[5]>0 else ''
            lines.append(f"{numbering} {b['title']}")
        payload = '\n'.join(lines)
        # show copy-to-clipboard widget
        components.html(f"<textarea id=\"t\" style=\"width:100%;height:200px;\">{payload}</textarea><br><button onclick=\"navigator.clipboard.writeText(document.getElementById('t').value)\">Copier dans le presse-papiers</button>",height=280)

st.write('---')
st.write('Petit mode d\'emploi: utilise + pour ajouter, clique sur un bloc pour le modifier, utilise les flèches pour réordonner, et le pad dans la barre latérale pour ajouter commentaires/pitch/resumé.')
